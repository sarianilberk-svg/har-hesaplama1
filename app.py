import streamlit as st
from docx import Document
from io import BytesIO
import re

# Sayfa Ayarları
st.set_page_config(page_title="UYAP Karar Asistanı V1.3", layout="wide")
st.title("⚖️ UYAP Karar Asistanı V1.3")
st.markdown("Hızlı Karar Hesaplama, **Arabuluculuk Modülü** ve **Word'e Aktarma** Özelliği")

# Türk Lirası Formatlayıcı
def tr_format(miktar):
    return f"{miktar:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

# AAÜT Vekalet Ücreti Hesaplama Fonksiyonu
def hesapla_vekalet(miktar):
    if miktar <= 0: return 0.0
    ucret = 0.0
    if miktar <= 600000: ucret = miktar * 0.16
    elif miktar <= 1200000: ucret = 96000 + (miktar - 600000) * 0.15
    elif miktar <= 2400000: ucret = 186000 + (miktar - 1200000) * 0.14
    elif miktar <= 4800000: ucret = 354000 + (miktar - 2400000) * 0.11
    elif miktar <= 9600000: ucret = 618000 + (miktar - 4800000) * 0.08
    elif miktar <= 19200000: ucret = 1002000 + (miktar - 9600000) * 0.05
    elif miktar <= 38400000: ucret = 1482000 + (miktar - 19200000) * 0.03
    else: ucret = 2058000 + (miktar - 38400000) * 0.01
        
    maktu_sinir = 45000.0
    if ucret < maktu_sinir:
        ucret = min(maktu_sinir, miktar)
    return ucret

# 1. BÖLÜM: Veri Girişi
st.header("📝 1. Dava Bilgileri ve Masraflar")

dava_turu = st.selectbox("Dava ve Harç Türünü Seçiniz", [
    "Genel Nispi (Asliye Hukuk, Ticaret vb. - Binde 68.31)",
    "Sulh Hukuk Davaları (Binde 13.65)",
    "Tahliye Davaları (Binde 11.68)",
    "Tüketici Mahkemesi (Davacı Tüketici / Harçtan Muaf)"
])

muafiyet = False
if "68.31" in dava_turu: oran = 68.31 / 1000; maktu = 732.0
elif "13.65" in dava_turu: oran = 13.65 / 1000; maktu = 732.0
elif "11.68" in dava_turu: oran = 11.68 / 1000; maktu = 732.0
elif "Tüketici" in dava_turu: oran = 68.31 / 1000; maktu = 0.0; muafiyet = True

col1, col2 = st.columns(2)

with col1:
    talep = st.number_input("Dava Değeri / Talep (TL)", min_value=0.0, value=100000.0, step=1000.0)
    kabul = st.number_input("Kabul Edilen Tutar (TL)", min_value=0.0, value=60000.0, step=1000.0)
    
    if muafiyet:
        pesin_harc = st.number_input("Yatan Peşin + Islah Harcı (TL) - (Davacı Muaf)", min_value=0.0, value=0.0, disabled=True)
    else:
        pesin_harc = st.number_input("Yatan Peşin + Islah Harcı (TL)", min_value=0.0, value=1707.75, step=10.0)
    
    st.subheader("Taraf Temsili ve Arabuluculuk")
    davaci_vekil = st.checkbox("Davacı kendisini vekille temsil ettirdi", value=True)
    arabuluculuk_var = st.checkbox("Dava Şartı Arabuluculuk Var", value=False)

    if kabul > talep:
        st.error("Kabul edilen tutar, dava değerinden büyük olamaz!")
        kabul = talep

with col2:
    if muafiyet:
        basvuru = st.number_input("Başvurma Harcı (TL) - (Davacı Muaf)", min_value=0.0, value=0.0, disabled=True)
    else:
        basvuru = st.number_input("Başvurma Harcı (TL)", min_value=0.0, value=732.0, step=10.0)
        
    d_gider = st.number_input("Davacı Harç Dışı Gideri (TL)", min_value=0.0, value=5000.0, step=100.0)
    dv_gider = st.number_input("Davalı Harç Dışı Gideri (TL)", min_value=0.0, value=0.0, step=100.0)
    
    st.subheader(" ") 
    davali_vekil = st.checkbox("Davalı kendisini vekille temsil ettirdi", value=False)
    
    if arabuluculuk_var:
        ara_ucret = st.number_input("Arabuluculuk Ücreti (TL)", min_value=0.0, value=3120.0, step=10.0)
    else:
        ara_ucret = 0.0

# 2. BÖLÜM: Hesaplamalar
reddedilen = talep - kabul
kabul_orani = kabul / talep if talep > 0 else 0
ret_orani = reddedilen / talep if talep > 0 else 0

alinmasi_gereken_harc = kabul * oran
if not muafiyet and alinmasi_gereken_harc < maktu:
    alinmasi_gereken_harc = maktu

bakiye_harc = alinmasi_gereken_harc - pesin_harc
iade_harc = 0.0
if bakiye_harc < 0:
    iade_harc = abs(bakiye_harc)
    bakiye_harc = 0.0

d_gider_davali_payi = d_gider * kabul_orani
dv_gider_davaci_payi = dv_gider * ret_orani
ara_davali_payi = ara_ucret * kabul_orani
ara_davaci_payi = ara_ucret * ret_orani

davaci_vekalet = hesapla_vekalet(kabul) if davaci_vekil else 0.0
davali_vekalet = hesapla_vekalet(reddedilen) if davali_vekil else 0.0

st.divider()

# 3. BÖLÜM: Hüküm Metni Oluşturma
hukum = "HÜKÜM: Gerekçesi ekli kararda açıklanacağı üzere;\n\n"

if kabul == talep:
    hukum += f"**1-** Davanın KABULÜNE, {tr_format(kabul)} TL'nin davalıdan alınarak davacıya VERİLMESİNE,\n\n"
elif kabul == 0:
    hukum += f"**1-** Davanın REDDİNE,\n\n"
else:
    hukum += f"**1-** Davanın KISMEN KABULÜNE, KISMEN REDDİNE;\n"
    hukum += f"    **a)** {tr_format(kabul)} TL'nin davalıdan alınarak davacıya VERİLMESİNE,\n"
    hukum += f"    **b)** Fazlaya ilişkin {tr_format(reddedilen)} TL'lik talebin REDDİNE,\n\n"

if muafiyet:
    if kabul > 0:
        hukum += f"**2-** Davacı taraf harçtan muaf olduğundan, kabul edilen miktar üzerinden hesaplanan {tr_format(alinmasi_gereken_harc)} TL karar ve ilam harcının DAVALIDAN ALINARAK HAZİNE'YE İRAT KAYDINA,\n\n"
    else:
        hukum += f"**2-** Davacı taraf harçtan muaf olduğundan başkaca harç alınmasına YER OLMADIĞINA,\n\n"
else:
    hukum += f"**2-** Karar tarihinde yürürlükte bulunan Harçlar Tarifesi gereğince, alınması gereken {tr_format(alinmasi_gereken_harc)} TL nispi karar ve ilam harcından, davacı tarafından peşin yatırılan {tr_format(pesin_harc)} TL harcın mahsubu ile bakiye {tr_format(bakiye_harc)} TL harcın DAVALIDAN ALINARAK HAZİNE'YE İRAT KAYDINA,\n\n"
    if iade_harc > 0:
        hukum += f"**-** Mahsup sonrası fazla yatırıldığı anlaşılan {tr_format(iade_harc)} TL harcın karar kesinleştiğinde istek halinde DAVACIYA İADESİNE,\n\n"

if not muafiyet:
    hukum += f"**3-** Davacı tarafından yatırıldığı anlaşılan {tr_format(pesin_harc)} TL karar-ilam harcı ile {tr_format(basvuru)} TL başvurma harcı toplamı olan {tr_format(pesin_harc + basvuru)} TL harcın DAVALIDAN ALINARAK DAVACIYA VERİLMESİNE,\n\n"
madde_no = 4 if not muafiyet else 3

if d_gider > 0:
    if kabul == talep: hukum += f"**{madde_no}-** Davacı tarafından yapılan harç dışı toplam {tr_format(d_gider)} TL yargılama giderinin DAVALIDAN ALINARAK DAVACIYA VERİLMESİNE,\n\n"
    elif kabul == 0: hukum += f"**{madde_no}-** Davacı tarafından yapılan harç dışı toplam {tr_format(d_gider)} TL yargılama giderinin davacı üzerinde BIRAKILMASINA,\n\n"
    else: hukum += f"**{madde_no}-** Davacı tarafından yapılan harç dışı toplam {tr_format(d_gider)} TL yargılama giderinden haklılık oranına isabet eden {tr_format(d_gider_davali_payi)} TL'nin DAVALIDAN ALINARAK DAVACIYA VERİLMESİNE, bakiye kısmın davacı üzerinde BIRAKILMASINA,\n\n"
    madde_no += 1

if dv_gider > 0:
    hukum += f"**{madde_no}-** Davalı tarafından yapılan harç dışı toplam {tr_format(dv_gider)} TL yargılama giderinden haklılık oranına isabet eden {tr_format(dv_gider_davaci_payi)} TL'nin DAVACIDAN ALINARAK DAVALIYA VERİLMESİNE, bakiye kısmın davalı üzerinde BIRAKILMASINA,\n\n"
else:
    hukum += f"**{madde_no}-** Davalı tarafından yapılan yargılama gideri bulunmadığından bu hususta karar verilmesine YER OLMADIĞINA,\n\n"
madde_no += 1

if davaci_vekil and kabul > 0:
    hukum += f"**{madde_no}-** Davacı yararına AAÜT gereğince hesaplanan {tr_format(davaci_vekalet)} TL vekalet ücretinin DAVALIDAN ALINARAK DAVACIYA VERİLMESİNE,\n\n"
    madde_no += 1

if davali_vekil and reddedilen > 0:
    hukum += f"**{madde_no}-** Davalı yararına AAÜT gereğince hesaplanan {tr_format(davali_vekalet)} TL vekalet ücretinin DAVACIDAN ALINARAK DAVALIYA VERİLMESİNE,\n\n"
    madde_no += 1

if arabuluculuk_var:
    if kabul == talep: hukum += f"**{madde_no}-** Dava şartı arabuluculuk faaliyeti nedeniyle Adalet Bakanlığı bütçesinden ödenen {tr_format(ara_ucret)} TL arabuluculuk ücretinin DAVALIDAN ALINARAK HAZİNE'YE İRAT KAYDINA,\n\n"
    elif kabul == 0: hukum += f"**{madde_no}-** Dava şartı arabuluculuk faaliyeti nedeniyle Adalet Bakanlığı bütçesinden ödenen {tr_format(ara_ucret)} TL arabuluculuk ücretinin DAVACIDAN ALINARAK HAZİNE'YE İRAT KAYDINA,\n\n"
    else: hukum += f"**{madde_no}-** Dava şartı arabuluculuk faaliyeti nedeniyle Adalet Bakanlığı bütçesinden ödenen {tr_format(ara_ucret)} TL arabuluculuk ücretinden haklılık oranına isabet eden {tr_format(ara_davaci_payi)} TL'nin davacıdan, {tr_format(ara_davali_payi)} TL'nin davalıdan ALINARAK HAZİNE'YE İRAT KAYDINA,\n\n"
    madde_no += 1

hukum += f"**{madde_no}-** Taraflarca yatırılan gider avansından kullanılmayan kısmın karar kesinleştiğinde HMK m.333 uyarınca yatırana İADESİNE karar verildi."

# Ekranda Gösterme ve İndirme Bölümü
st.header("📜 2. Otomatik Hüküm Fıkrası")

# Word Çıktısı Üreten Fonksiyon
def create_word_doc(text):
    doc = Document()
    doc.add_heading('HÜKÜM', 1)
    # Metindeki Markdown işaretlerini temizle
    clean_text = re.sub(r'\*\*', '', text)
    clean_text = clean_text.replace("HÜKÜM: Gerekçesi ekli kararda açıklanacağı üzere;\n\n", "")
    
    for line in clean_text.split('\n'):
        if line.strip():
            doc.add_paragraph(line.strip())
            
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

word_file = create_word_doc(hukum)

col_out1, col_out2 = st.columns([3, 1])
with col_out1:
    st.info(hukum)
with col_out2:
    st.download_button(
        label="📄 Word Olarak İndir",
        data=word_file,
        file_name="Hukum_Fikrasi.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )
